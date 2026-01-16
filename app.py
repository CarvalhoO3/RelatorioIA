import sys
import time
import openpyxl
import os
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from tkinter import filedialog, Tk


client = OpenAI(api_key="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")  # Substitua pela sua chave da OpenAI

MODELO_LLM = "gpt-4.1-mini"

# ✅ Seletor de arquivos para escolher manualmente
Tk().withdraw()  # Oculta a janela principal do Tkinter
print("📂 Selecione a planilha Excel...")
CAMINHO_PLANILHA = filedialog.askopenfilename(title="Selecione a planilha", filetypes=[("Excel files", "*.xlsx")])

print("📂 Selecione o modelo Word...")
CAMINHO_MODELO_WORD = filedialog.askopenfilename(title="Selecione o modelo Word", filetypes=[("Word files", "*.docx")])

def call_api_with_retry(prompt, max_retries=5):
    for i in range(max_retries):
        try:
            print(f"🤖 Gerando texto (tentativa {i + 1})...")
            response = client.chat.completions.create(
                model=MODELO_LLM,
                messages=[
                    {"role": "system", "content": "Você é um assistente útil que gera relatórios técnicos."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip()

        except Exception as e:
            wait = 2 ** i
            print(f"⚠ Erro OpenAI: {e} | aguardando {wait}s")
            time.sleep(wait)

    return "ERRO AO GERAR RELATÓRIO."

def parse_markdown_report(texto):
    secoes = {
        "Participantes": "",
        "Introdução": "",
        "Detalhes da Visita": "",
        "Conclusão": ""
    }

    atual = None
    for line in texto.splitlines():
        line_strip = line.strip()

        if line_strip.startswith("##"):
            titulo = line_strip.replace("#", "").strip()
            if titulo in secoes:
                atual = titulo
            else:
                atual = None
            continue

        if atual:
            secoes[atual] += line + "\n"

    return secoes

def preencher_modelo_word(secoes, nome_saida):
    if not os.path.exists(CAMINHO_MODELO_WORD):
        print(f"❌ Modelo Word não encontrado: {CAMINHO_MODELO_WORD}")
        return

    doc = Document(CAMINHO_MODELO_WORD)

    mapa = {
        "Participantes:": secoes["Participantes"],
        "Introdução:": secoes["Introdução"],
        "Detalhes da Visita:": secoes["Detalhes da Visita"],
        "Conclusão:": secoes["Conclusão"]
    }

    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto in mapa:
            p.text = ""  # ✅ Corrigido: limpa o parágrafo
            run = p.add_run(texto + "\n" + mapa[texto])
            run.font.size = Pt(11)

    doc.save(nome_saida)
    print(f"📄 Relatório gerado → {nome_saida}")

def processar_planilha(caminho):
    print("📖 Lendo planilha...")

    if not os.path.exists(caminho):
        print(f"❌ Planilha não encontrada: {caminho}")
        return

    wb = openpyxl.load_workbook(caminho)
    ws = wb.active

    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or len(row) < 4:  # ✅ Corrigido
            continue

        id_rel, participantes, introducao, detalhes = row[:4]

        if not id_rel:
            continue

        print(f"\n➡ Processando relatório ID {id_rel}")

        prompt =  f"""
Você é um engenehiro técnico e voltou de uma visita técnica. Com base nas informações abaixo, gere um relatório técnico detalhado seguindo a estrutura obrigatória.

Utilize os textos fornecidos como BASE e:
- complemente com linguagem técnica
- expanda as ideias
- acrescente contexto profissional
- NÃO remova informações existentes

Estrutura obrigatória:

## Participantes
Liste e descreva brevemente os participantes, mantendo os nomes fornecidos e os cargos.
Texto base:
{participantes}

## Introdução
Reescreva e complemente a introdução, deixando-a mais técnica e formal.
Texto base:
{introducao}

## Detalhes da Visita
Amplie os detalhes técnicos da visita, mantendo as informações originais.
Texto base:
{detalhes}

## Conclusão
Crie uma conclusão técnica com base nos dados acima.

"""

        texto = call_api_with_retry(prompt)
        secoes = parse_markdown_report(texto)

        nome_saida = f"Relatorio_{id_rel}.docx"
        preencher_modelo_word(secoes, nome_saida)

def main():
    processar_planilha(CAMINHO_PLANILHA)
    print("\n✔ Processo finalizado com sucesso!")
    input("ENTER para sair...")

if __name__ == "__main__":
    main()
