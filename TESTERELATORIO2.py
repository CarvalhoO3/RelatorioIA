import os
from groq import Groq
from docx import Document
from tkinter import filedialog, Tk, messagebox



def extrair_texto_word(caminho_doc):
    doc = Document(caminho_doc)
    
    textos = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(textos)

def extrair_secao(texto, inicio, fim=None):
   
    texto_upper = texto.upper()
    inicio_upper = inicio.upper()
    
    if inicio_upper not in texto_upper: 
        return "Informação não localizada no relato original."
    
    pos_inicio = texto_upper.find(inicio_upper) + len(inicio_upper)
    parte = texto[pos_inicio:]
    
    if fim:
        fim_upper = fim.upper()
        if fim_upper in parte.upper():
            pos_fim = parte.upper().find(fim_upper)
            parte = parte[:pos_fim]
            
    return parte.strip().replace("*", "") # Remove possíveis asteriscos de Markdown

# =========================
# EXECUÇÃO PRINCIPAL
# =========================

def executar():
    root = Tk()
    root.withdraw()

    # --- COLOQUE SUA CHAVE GROQ ABAIXO ---
    api_key = "SUA_CHAVE_GROQ_AQUI" 
    # -------------------------------------

    if not api_key or api_key == "SUA_CHAVE_GROQ_AQUI":
        messagebox.showerror("Erro", "Configure sua chave API da Groq no código.")
        return

    try:
        client = Groq(api_key=api_key)

        # Seleção de arquivos com avisos claros
        messagebox.showinfo("Passo 1", "Selecione o arquivo de DADOS (Ex: testerelat.docx)")
        origem = filedialog.askopenfilename(title="Relato Bruto", filetypes=[("Word", "*.docx")])
        if not origem: return
        
        messagebox.showinfo("Passo 2", "Selecione o MODELO (Ex: Modelo_padrão.docx)")
        modelo_path = filedialog.askopenfilename(title="Modelo Padrão", filetypes=[("Word", "*.docx")])
        if not modelo_path: return

        texto_bruto = extrair_texto_word(origem)

        # Prompt ultra-específico para evitar que a IA ignore os dados
        prompt = f"""
        TRANSFORME O RELATO ABAIXO EM UM RELATÓRIO TÉCNICO FORMAL.
        
        DADOS REAIS DISPONÍVEIS:
        {texto_bruto}

        REGRAS:
        1. Use apenas os fatos acima (Participantes, local, animais e refeição).
        2. Não adicione informações que não estão no texto.
        3. Use obrigatoriamente os cabeçalhos [PARTICIPANTES], [INTRODUCAO], [DETALHES] e [CONCLUSAO].

        RESPOSTA NO FORMATO:
        [PARTICIPANTES]
        (insira aqui os nomes encontrados)
        [INTRODUCAO]
        (contexto formal da visita)
        [DETALHES]
        (atividades realizadas)
        [CONCLUSAO]
        (encerramento técnico)
        """

        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.1 # Menor criatividade para evitar que a IA ignore os dados
        )

        texto_ia = chat_completion.choices[0].message.content

        # Mapeamento para os títulos do seu Modelo_padrão.docx [cite: 4, 5, 6, 7]
        dados_ia = {
            "Participantes:": extrair_secao(texto_ia, "[PARTICIPANTES]", "[INTRODUCAO]"),
            "Introdução:": extrair_secao(texto_ia, "[INTRODUCAO]", "[DETALHES]"),
            "Detalhes da Visita:": extrair_secao(texto_ia, "[DETALHES]", "[CONCLUSAO]"),
            "Conclusão:": extrair_secao(texto_ia, "[CONCLUSAO]")
        }

        doc = Document(modelo_path)
        
        # Percorre parágrafos e tabelas (caso as tags estejam em tabelas)
        for p in doc.paragraphs:
            for termo, conteudo in dados_ia.items():
                if termo in p.text:
                    p.add_run(f" {conteudo}")

        # Salva o arquivo final
        diretorio = os.path.dirname(origem)
        caminho_final = os.path.join(diretorio, "Relatorio_Finalizado_OK.docx")
        doc.save(caminho_final)

        messagebox.showinfo("Sucesso!", f"Relatório gerado com os dados reais!\n\nSalvo em:\n{caminho_final}")
        os.startfile(diretorio)

    except Exception as e:
        messagebox.showerror("Erro Crítico", f"Falha no processamento: {str(e)}")

if __name__ == "__main__":
    executar()